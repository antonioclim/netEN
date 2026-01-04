"""Generate a Week 7 DOCX lab sheet.

This is optional. The main kit works without it.
"""

from __future__ import annotations

import argparse
from pathlib import Path


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Generate the Week 7 lab sheet as DOCX.")
    p.add_argument("--out", default="docs/Week7_Lab_Sheet.docx", help="Output DOCX path")
    return p


def main() -> int:
    args = build_parser().parse_args()
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document  # type: ignore
    except Exception:
        print("python-docx is not installed. Install it with: python3 -m pip install python-docx")
        return 2

    doc = Document()
    doc.add_heading("Week 7 — Packet capture and filtering (Networking laboratory)", level=1)

    doc.add_paragraph(
        "This laboratory focuses on capturing TCP and UDP traffic and applying simple, reproducible filtering rules "
        "in a controlled Mininet environment."
    )

    doc.add_heading("Learning outcomes", level=2)
    for item in [
        "capture and inspect TCP and UDP traffic using tcpdump and tshark",
        "apply filtering rules and verify their effect with repeatable tests",
        "produce evidence in the form of pcaps, logs and short written conclusions",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Quickstart", level=2)
    doc.add_paragraph("From the kit root directory:")
    doc.add_paragraph("chmod +x scripts/*.sh", style="List Bullet")
    doc.add_paragraph("./scripts/setup.sh", style="List Bullet")
    doc.add_paragraph("./scripts/run_all.sh", style="List Bullet")
    doc.add_paragraph("./tests/smoke_test.sh", style="List Bullet")

    doc.add_heading("Student deliverable", level=2)
    doc.add_paragraph(
        "Create a new firewall profile that blocks one traffic class while leaving the rest functional. "
        "Provide a pcap, a log and a short explanation based on packet fields."
    )

    doc.add_heading("Evidence checklist", level=2)
    for item in [
        "artifacts/demo.pcap (or an equivalent capture you produced)",
        "artifacts/demo.log",
        "artifacts/validation.txt",
        "3–6 tshark lines that support your conclusion",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(out_path))
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
